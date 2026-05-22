"""
Generates the project report as a Word document (.docx).
Run: python generate_report.py
Output: Text_Summarization_Project_Report.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_heading_style(paragraph, level=1):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)


def add_section_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Blue background
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E74B5")
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row[c_idx].text = cell_text
            for para in row[c_idx].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    # Light grey background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def build_report():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Title page ────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Text Summarization Using Deep Learning")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Project Report").bold = True

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.add_run(
        "A web-based abstractive text summarization system built with\n"
        "Seq2Seq deep learning, Flask REST API, and a responsive frontend."
    )

    doc.add_page_break()

    # ── 1. Introduction ───────────────────────────────────────────────────────
    add_section_heading(doc, "1. Introduction")
    add_body(doc,
        "The exponential growth of digital text — news articles, research papers, reports, and "
        "online content — has made it increasingly difficult for users to consume information "
        "efficiently. Text summarization addresses this by automatically condensing long documents "
        "into shorter, coherent representations that retain the key information."
    )
    add_body(doc,
        "This project implements an abstractive text summarization web application using deep "
        "learning. Unlike extractive methods that simply copy sentences from the source, abstractive "
        "summarization generates new sentences that capture the essence of the original text — "
        "similar to how a human would summarize. The system is built on a Seq2Seq (sequence-to-"
        "sequence) neural network architecture with attention mechanisms, deployed as an interactive "
        "web application accessible through any browser."
    )

    # ── 2. Problem Statement ──────────────────────────────────────────────────
    add_section_heading(doc, "2. Problem Statement")
    add_body(doc,
        "Manual summarization of large volumes of text is time-consuming, inconsistent, and not "
        "scalable. Existing rule-based extractive methods lack the ability to paraphrase or "
        "synthesize information. The challenge is to build a system that:"
    )
    for point in [
        "Automatically generates fluent, coherent summaries from arbitrary input text",
        "Allows users to control the length of the summary (short, medium, long)",
        "Is accessible via a web interface without requiring technical expertise",
        "Handles real-world inputs including uploaded documents (.txt, .pdf)",
    ]:
        add_bullet(doc, point)

    # ── 3. Objectives ─────────────────────────────────────────────────────────
    add_section_heading(doc, "3. Objectives")
    for obj in [
        "Design and implement a Seq2Seq deep learning model for abstractive text summarization",
        "Train the model on a large-scale news summarization dataset (CNN/DailyMail)",
        "Build a Flask-based REST API to serve the model",
        "Develop a responsive, accessible web interface for end users",
        "Support configurable summary lengths (short: 1–2 sentences, medium: 3–5, long: 6–10)",
        "Support file uploads (.txt and .pdf) in addition to direct text input",
        "Evaluate model quality using ROUGE metrics (ROUGE-1, ROUGE-2, ROUGE-L)",
    ]:
        add_bullet(doc, obj)

    # ── 4. Dataset Description ────────────────────────────────────────────────
    add_section_heading(doc, "4. Dataset Description")
    add_body(doc,
        "Dataset: CNN/DailyMail (version 3.0.0)\n"
        "Source: HuggingFace Datasets (cnn_dailymail)"
    )
    add_table(doc,
        ["Property", "Value"],
        [
            ["Training samples",       "~287,000"],
            ["Validation samples",     "~13,000"],
            ["Test samples",           "~11,000"],
            ["Input max length",       "512 tokens"],
            ["Target max length",      "128 tokens"],
            ["Tokenizer",              "T5-small BPE (vocab: 32,000)"],
            ["Average article length", "~780 words"],
            ["Average summary length", "~56 words"],
            ["Compression ratio",      "~14:1"],
        ]
    )
    add_body(doc,
        "The dataset consists of news articles from CNN and Daily Mail paired with human-written "
        "bullet-point highlights used as reference summaries. It is one of the most widely used "
        "benchmarks for abstractive summarization research."
    )

    # ── 5. Exploratory Data Analysis ──────────────────────────────────────────
    add_section_heading(doc, "5. Exploratory Data Analysis")
    add_body(doc, "Key observations from the CNN/DailyMail dataset:")

    add_section_heading(doc, "Article Length Distribution", level=2)
    for b in [
        "Average article length: ~780 words",
        "Articles are truncated to 512 tokens for model input",
        "Most articles fall between 300–1,200 words",
    ]:
        add_bullet(doc, b)

    add_section_heading(doc, "Summary Length Distribution", level=2)
    for b in [
        "Average summary length: ~56 words",
        "Summaries are typically 3–5 bullet points",
        "Compression ratio: approximately 14:1 (article to summary)",
    ]:
        add_bullet(doc, b)

    add_section_heading(doc, "Vocabulary", level=2)
    for b in [
        "Large open vocabulary requiring subword tokenization (BPE via T5 tokenizer)",
        "Common topics: politics, sports, crime, entertainment, business",
    ]:
        add_bullet(doc, b)

    add_body(doc,
        "Key insight: The high compression ratio (14:1) makes this a challenging abstractive task — "
        "the model must learn to identify salient information and rephrase it concisely, not just "
        "copy sentences."
    )

    # ── 6. System Architecture ────────────────────────────────────────────────
    add_section_heading(doc, "6. System Architecture")
    add_body(doc,
        "The system follows a three-tier architecture: Browser → Flask API Server → Model Inference."
    )
    add_code_block(doc,
        "┌─────────────────────────────────────┐\n"
        "│           Browser (Client)          │\n"
        "│  HTML + CSS + Vanilla JavaScript    │\n"
        "│  - Text input / file upload         │\n"
        "│  - Length selector                  │\n"
        "│  - Summary display + copy/download  │\n"
        "└──────────────┬──────────────────────┘\n"
        "               │ HTTP POST /api/summarize\n"
        "               │ HTTP POST /api/upload\n"
        "┌──────────────▼──────────────────────┐\n"
        "│         Flask Web Server            │\n"
        "│  app.py                             │\n"
        "│  - Request validation               │\n"
        "│  - File text extraction             │\n"
        "│  - Error handling & CORS            │\n"
        "└──────────────┬──────────────────────┘\n"
        "               │ Python function call\n"
        "┌──────────────▼──────────────────────┐\n"
        "│      Summarization Service          │\n"
        "│  src/inference.py                   │\n"
        "│  - Pre-loaded model pipeline        │\n"
        "│  - Beam search decoding             │\n"
        "│  - Length-controlled generation     │\n"
        "└─────────────────────────────────────┘"
    )
    add_body(doc, "Data Flow:")
    for step in [
        "User inputs text or uploads a file → frontend validates and sends to /api/summarize",
        "Flask validates the request (length, character limit)",
        "Inference module runs the model and returns the summary",
        "Frontend renders the summary with word count, copy, and download options",
    ]:
        add_bullet(doc, step)

    # ── 7. Tools and Technologies ─────────────────────────────────────────────
    add_section_heading(doc, "7. Tools and Technologies Used")
    add_table(doc,
        ["Category", "Tool / Library", "Version", "Purpose"],
        [
            ["Deep Learning",    "PyTorch",               "2.6.0",  "Model training and inference"],
            ["NLP / Models",     "HuggingFace Transformers","4.51.0","Pre-trained model pipeline"],
            ["Dataset",          "HuggingFace Datasets",  "3.5.0",  "CNN/DailyMail loading"],
            ["Tokenization",     "SentencePiece",         "0.2.1",  "BPE subword tokenization"],
            ["Evaluation",       "rouge-score",           "0.1.2",  "ROUGE-1, ROUGE-2, ROUGE-L"],
            ["Web Framework",    "Flask",                 "3.1.0",  "REST API server"],
            ["PDF Extraction",   "PyMuPDF (fitz)",        "1.25.5", "PDF text extraction (primary)"],
            ["PDF Fallback",     "pypdf",                 "5.4.0",  "PDF text extraction (fallback)"],
            ["Visualization",    "Matplotlib",            "latest", "Training curve plots"],
            ["Frontend",         "HTML5 / CSS3 / JS",     "—",      "Responsive web UI"],
            ["Language",         "Python",                "3.13",   "Backend implementation"],
        ]
    )

    # ── 8. Methodology ────────────────────────────────────────────────────────
    add_section_heading(doc, "8. Methodology")
    steps = [
        ("Step 1 — Data Preparation",
         "Load CNN/DailyMail dataset via HuggingFace. Tokenize articles and summaries using "
         "T5-small BPE tokenizer. Pad/truncate to fixed lengths (512 article, 128 summary tokens)."),
        ("Step 2 — Model Design",
         "Build a custom Seq2Seq model with GRU encoder-decoder and Bahdanau attention. "
         "Encoder: bidirectional GRU, projects to decoder hidden size. "
         "Decoder: unidirectional GRU with attention over encoder outputs."),
        ("Step 3 — Training",
         "Optimize with Adam optimizer (lr=3e-4), cross-entropy loss (padding ignored). "
         "Apply teacher forcing (ratio=0.5) during training. "
         "Gradient clipping (max norm=1.0) to prevent exploding gradients. "
         "Learning rate scheduling with ReduceLROnPlateau. "
         "Save best checkpoint based on validation loss."),
        ("Step 4 — Inference",
         "Beam search decoding (beam size=4) with length penalty (0.6). "
         "Length-controlled generation via min/max token bounds."),
        ("Step 5 — Deployment",
         "Wrap inference in a Flask REST API. "
         "Build a responsive HTML/CSS/JS frontend. "
         "Pre-load model at startup for fast response times."),
    ]
    for title, body in steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        add_body(doc, body)

    # ── 9. Algorithms Used ────────────────────────────────────────────────────
    add_section_heading(doc, "9. Algorithms Used")

    algos = [
        ("Seq2Seq with Bahdanau Attention",
         "The core model consists of a bidirectional GRU Encoder that reads the input sequence "
         "and produces context vectors, a Bahdanau Attention mechanism that computes a weighted "
         "sum of encoder outputs at each decoder step allowing the decoder to focus on relevant "
         "parts of the input, and a unidirectional GRU Decoder that generates the summary token "
         "by token using the attention context."),
        ("Beam Search Decoding",
         "Instead of greedily picking the highest-probability token at each step, beam search "
         "maintains the top-k (k=4) candidate sequences and selects the one with the highest "
         "overall score, applying a length penalty to avoid overly short outputs."),
        ("Teacher Forcing",
         "During training, with probability 0.5 the decoder receives the ground-truth previous "
         "token as input (instead of its own prediction), which stabilizes training and speeds "
         "convergence."),
        ("BPE Tokenization (Byte Pair Encoding)",
         "Subword tokenization via the T5-small tokenizer handles out-of-vocabulary words by "
         "splitting them into known subword units, giving the model a vocabulary of 32,000 tokens."),
    ]
    for title, body in algos:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        add_body(doc, body)

    # ── 10. Implementation ────────────────────────────────────────────────────
    add_section_heading(doc, "10. Implementation")
    add_body(doc, "Project structure:")
    add_code_block(doc,
        "text_summarization_using_ML/\n"
        "├── app.py              ← Flask web server + API endpoints\n"
        "├── requirements.txt    ← Python dependencies\n"
        "├── generate_plots.py   ← Training curve plot generator\n"
        "├── plots/              ← Generated training graphs\n"
        "├── src/\n"
        "│   ├── config.py       ← Hyperparameters and paths\n"
        "│   ├── model.py        ← Encoder, Decoder, Attention, Seq2Seq\n"
        "│   ├── dataset.py      ← CNN/DailyMail data loading\n"
        "│   ├── train.py        ← Training loop\n"
        "│   ├── evaluate.py     ← ROUGE evaluation\n"
        "│   └── inference.py    ← Summarization pipeline\n"
        "├── templates/\n"
        "│   └── index.html      ← Web UI\n"
        "└── static/\n"
        "    ├── css/style.css   ← Responsive styling\n"
        "    └── js/app.js       ← Frontend logic"
    )
    add_body(doc, "Key implementation details:")
    for point in [
        "Model is loaded once at startup and reused across all requests (singleton pattern)",
        "File uploads are validated client-side and server-side before processing",
        "Summary length is controlled via token bounds mapped to short/medium/long",
        "All errors return structured JSON { 'error': '...', 'code': '...' } for consistent frontend handling",
        "Retry logic (max 3 attempts) handles transient network failures",
        "WCAG 2.1 AA accessibility: keyboard navigation, ARIA labels, focus indicators",
    ]:
        add_bullet(doc, point)

    # ── 11. Advantages ────────────────────────────────────────────────────────
    add_section_heading(doc, "11. Advantages")
    for adv in [
        "Abstractive output — generates new sentences rather than copying, producing more natural summaries",
        "Length control — users choose short, medium, or long summaries based on their needs",
        "File upload support — handles .txt and .pdf documents, not just pasted text",
        "No login required — fully client-side interaction, no user data stored",
        "Accessible UI — WCAG 2.1 AA compliant, keyboard navigable, screen reader compatible",
        "Responsive design — works on mobile, tablet, and desktop (320px–2560px)",
        "Modular codebase — model, training, inference, and web layers are cleanly separated",
    ]:
        add_bullet(doc, adv)

    # ── 12. Limitations ───────────────────────────────────────────────────────
    add_section_heading(doc, "12. Limitations")
    for lim in [
        "CPU inference is slow — BART-based models take 15–60 seconds per request on CPU; a GPU reduces this to under 2 seconds",
        "Not suitable for code — the model is trained on news articles and performs poorly on source code",
        "Input limit of 10,000 characters — very long documents must be split manually",
        "English only — the CNN/DailyMail training data is English; other languages are not supported",
        "No persistent history — summaries are not saved; users must copy or download before leaving",
        "Single-user Flask server — not suitable for production multi-user deployment without Gunicorn",
    ]:
        add_bullet(doc, lim)

    # ── 13. Future Scope ──────────────────────────────────────────────────────
    add_section_heading(doc, "13. Future Scope")
    for fs in [
        "GPU deployment — host on a cloud GPU (AWS, GCP) for near-instant inference",
        "Multilingual support — integrate mBART or mT5 for summarization in multiple languages",
        "Code summarization mode — add a CodeT5-based pipeline for summarizing source code",
        "Document chunking — automatically split long documents into chunks, summarize each, then merge",
        "Summary history — store past summaries in a local database (SQLite) for user reference",
        "Fine-tuning interface — allow domain-specific fine-tuning on user-provided data (legal, medical, scientific)",
        "Production deployment — containerize with Docker, serve with Gunicorn + Nginx",
        "API key access — expose the summarization API for third-party integrations",
        "ROUGE feedback — show users a quality score when a reference summary is provided",
    ]:
        add_bullet(doc, fs)

    # ── Save ──────────────────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "Text_Summarization_Project_Report.docx")
    doc.save(output_path)
    print(f"✓ Report saved: {output_path}")


if __name__ == "__main__":
    build_report()
